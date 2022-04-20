import os
import glob
import re
import shutil  # сохранение файла
import sys
from copy import deepcopy  # копирование таблиц

import requests  # загрузка файла из сети
from docx import Document  # работа с docx
from docx.shared import Cm, Pt, RGBColor
from kinopoisk_unofficial.kinopoisk_api_client import KinopoiskApiClient
from kinopoisk_unofficial.request.films.film_request import FilmRequest
from kinopoisk_unofficial.request.staff.staff_request import StaffRequest
from kinopoisk.movie import Movie  # api для поиска фильмов
from mutagen.mp4 import MP4, MP4Cover  # работа тегами
from PIL import Image  # работа с изображениями
from rich.panel import Panel
from rich.console import Console
from rich.columns import Columns

import config

ver = '0.4.5'
api = config.api_key
console = Console()


def is_api_ok(api):
    '''Проверка авторизации.'''
    try:
        api_client = KinopoiskApiClient(api)
        request = FilmRequest(328)
        response = api_client.films.send_film_request(request)
    except:
        return False
    else:
        return True


def get_film_info(film_code, api):
    '''
    Получение информации о фильме с помощью kinopoisk_api_client.

            Элементы списка:            
                0 - название фильма на русском языке
                1 - год
                2 - рейтинг Кинопоиска
                3 - список стран
                4 - описание
                5 - ссылка на постер
                6 - имя файла без расширения
                7 - режиссер
             8:17 - 10 актеров
    '''
    api_client = KinopoiskApiClient(api)
    request_staff = StaffRequest(film_code)
    response_staff = api_client.staff.send_staff_request(request_staff)
    stafflist = []
    for i in range(0, 11):  # загружаем 11 персоналий (режиссер + 10 актеров)
        if response_staff.items[i].name_ru == '':
            stafflist.append(response_staff.items[i].name_en)
        else:
            stafflist.append(response_staff.items[i].name_ru)

    request_film = FilmRequest(film_code)
    response_film = api_client.films.send_film_request(request_film)

    # с помощью регулярного выражения находим значение стран в кавычках ''
    countries = re.findall("'([^']*)'", str(response_film.film.countries))

    # имя файла
    filename = response_film.film.name_ru
    # очистка имени файла от запрещенных символов
    trtable = filename.maketrans('', '', '\/:*?"<>')
    filename = filename.translate(trtable)
    filmlist = [
        response_film.film.name_ru, response_film.film.year, response_film.film.rating_kinopoisk, countries,
        response_film.film.description, response_film.film.poster_url, filename
    ]
    return filmlist + stafflist


def write_film_to_table(current_table, filminfo):
    '''Заполнение таблицы в файле docx.'''
    paragraph = current_table.cell(0, 1).paragraphs[0]  # название фильма + рейтинг
    run = paragraph.add_run(str(filminfo[0]) + ' - ' + 'Кинопоиск ' + str(filminfo[2]))
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True

    paragraph = current_table.cell(1, 1).add_paragraph()  # год
    run = paragraph.add_run(str(filminfo[1]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # страна
    run = paragraph.add_run(', '.join(filminfo[3]))
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()  # режиссер
    run = paragraph.add_run('Режиссер: ' + filminfo[7])
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    paragraph = current_table.cell(1, 1).add_paragraph()

    paragraph = current_table.cell(1, 1).add_paragraph()  # в главных ролях
    run = paragraph.add_run('В главных ролях: ')
    run.font.color.rgb = RGBColor(255, 102, 0)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run = paragraph.add_run(', '.join(filminfo[8:]))
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.underline = True

    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()
    paragraph = current_table.cell(1, 1).add_paragraph()  # синопсис
    run = paragraph.add_run(filminfo[4])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    paragraph = current_table.cell(1, 1).add_paragraph()

    # загрузка постера
    image_url = filminfo[5]
    if not os.path.isdir("covers"):
        os.mkdir("covers")
    file_path = './covers/' + str(filminfo[6] + '.jpg')
    resp = requests.get(image_url, stream=True)
    if resp.status_code == 200:
        resp.raw.decode_content = True
        with open(file_path, 'wb') as f:  # открываем файл для бинарной записи
            shutil.copyfileobj(resp.raw, f)
    else:
        console.print(f'Не удалось загрузить постер ({image_url})')

    # изменение размера постера
    image = Image.open(file_path)
    width, height = image.size
    # обрезка до соотношения сторон 1x1.5
    if width > (height / 1.5):
        image = image.crop((((width - height / 1.5) / 2), 0, ((width - height / 1.5) / 2) + height / 1.5, height))
    image.thumbnail((360, 540))
    rgb_image = image.convert('RGB')  # для исправление возможной ошибки "OSError: cannot write mode RGBA as JPEG"
    rgb_image.save(file_path)

    # запись постера в таблицу
    paragraph = current_table.cell(0, 0).paragraphs[1]
    run = paragraph.add_run()
    run.add_picture(file_path, width=Cm(7))


def copy_table_after(table, paragraph):
    '''Копирование таблицы в указанный параграф.'''
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)


def clone_first_table(document: Document, num):
    '''Клонирует первую таблицу в документе num раз.'''
    template = document.tables[0]
    paragraph = document.paragraphs[0]
    for i in range(num):
        copy_table_after(template, paragraph)
        paragraph = document.add_paragraph()


def write_tags_to_mp4(film):
    '''Запись тегов в файл mp4.'''
    file_path = str(film[6] + '.mp4')
    if not os.path.isfile(file_path):
        console.print(f'Ошибка: Файл "{file_path}" не найден!')
        return
    video = MP4(file_path)
    video.delete()  # удаление всех тегов
    video["\xa9nam"] = film[0]  # title
    video["desc"] = film[4]  # description
    video["ldes"] = film[4]  # long description
    video["\xa9day"] = str(film[1])  # year
    cover = './covers/' + str(film[6] + '.jpg')
    with open(cover, "rb") as f:
        video["covr"] = [MP4Cover(f.read(), imageformat=MP4Cover.FORMAT_JPEG)]
    video.save()
    console.print(f'{file_path} - [bright_green]тег записан[/bright_green]')


def get_resource_path(relative_path):
    '''
    Определение пути для запуска из автономного exe файла.

    Pyinstaller cоздает временную папку, путь в _MEIPASS.
    '''
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def input_kinopoisk_id(choice):
    '''Функция поиска kinopoisk_id фильмов несколькими способами.'''
    if choice == 1:
        filmsearch = []
        filmlistprint = []
        while True:
            if len(filmlistprint) > 0:
                console.print("В списке следующие фильмы:")
                films_renderables = [Panel(str(film)) for film in filmlistprint]
                console.print(Columns(films_renderables))
            search = console.input('Введите название фильма и год выпуска или [b]Enter[/b] чтобы продолжить: ')
            if search == '':
                return filmsearch
            try:
                movie_list = Movie.objects.search(search)
            except Exception:
                console.print('[red]Фильм не найден, возникла ошибка!')
                continue
            else:
                if len(movie_list) < 1:
                    console.print('Фильм не найден.')
                    continue
                id = str(movie_list[0].id)
                console.print(f'[white bold]{movie_list[0]} (kinopoisk_id: {id})')
                # console.print(f"Kinopoisk_id: {id}")
                choice_1 = console.input(
                    'Варианты: Добавить в список ([b]1[/b]), новый поиск ([b]2[/b]), закончить и продолжить ([b]Enter[/b]): '
                )
                if choice_1 == '1':
                    filmsearch.append(id)
                    filmlistprint.append(f"{movie_list[0].title}, {movie_list[0].year}")
                elif choice_1 == '2':
                    continue
                elif choice_1 == '':
                    return filmsearch
    elif choice == 2:
        inputstr = console.input('Введите через пробел идентификаторы фильмов ([b]kinopoisk id[/b]): ')
        return inputstr.split()
    elif choice == 3:
        filmsearch = []
        mp4files = glob.glob('*.mp4')
        if len(mp4files) < 1:
            return filmsearch
        for file in mp4files:
            search = file[:-4]
            console.print(f'Поиск: {search}...')
            try:
                movie_list = Movie.objects.search(search)
            except Exception:
                console.print('[red]Фильм не найден, возникла ошибка!')
                continue
            else:
                if len(movie_list) < 1:
                    console.print('Фильм не найден.')
                    continue
                id = str(movie_list[0].id)
                console.print(f'Найден фильм: [orange]{movie_list[0]}[/orange], kinopoisk id: {id}')
                filmsearch.append(id)
        return filmsearch


def clean_and_exit():
    '''
    Очищает каталог по запросу, встает на паузу и выходит.
    
    Удаляет каталог covers и файл list.txt.
    '''
    ask = str(console.input('Произвести очистку каталога (папка covers, list.txt)? [white bold](y,1/n,2) '))
    if ask.lower() in ('y', '1', 'д'):
        os.remove("list.txt")
        shutil.rmtree("./covers/")
        console.print("Каталог очищен.")
    os.system('pause')
    sys.exit()


def pause_and_exit(*phrases:str):
    '''Выводит фразы из аргументов, встает на паузу и выходит.'''
    for phrase in phrases:
        console.print(phrase)
    os.system('pause')
    sys.exit()


terminal_size = os.get_terminal_size().columns - 1
console.print(
    Panel("Kinolist: Программа создания списка фильмов".center(terminal_size, " ") + '\n' +
          ver.center(terminal_size, " ")))

if not is_api_ok(api):
    console.print('[red]Ошибка API!')
    pause_and_exit()

file_path = get_resource_path('template.docx')  # определяем путь до шаблона
try:
    doc = Document(file_path)  # открываем шаблон
except Exception:
    pause_and_exit('[red]Ошибка! Не найден шаблон "template.docx". Список не создан.', '', 'Работа программы завершена.')

# считываем значения из файла list.txt
film_codes = []
if os.path.isfile('./list.txt'):
    file_list = open('./list.txt', 'r')
    lines = file_list.readlines()  # считываем все строки
    file_list.close()
    for line in lines:
        film_codes.append(line.strip())
    if len(film_codes) < 1:
        pause_and_exit('В списке 0 фильмов. Работа программы завершена.')
    console.print(f'Найден файл "list.txt" (записей: {len(film_codes)})')
else:
    console.print('Файл "list.txt" не найден!')
    while True:
        choice = console.input(
            'Выберите режим: Поиск фильмов по названию ([b]1[/b]); ручной ввод kinopoisk_id ([b]2[/b]); поиск по mp4 файлам ([b]3[/b]); [b]Enter[/b] чтобы выйти: '
        )
        if choice == "1":
            film_codes = input_kinopoisk_id(1)
            break
        elif choice == "2":
            film_codes = input_kinopoisk_id(2)
            break
        elif choice == "3":
            film_codes = input_kinopoisk_id(3)
            break
        elif choice == "":
            pause_and_exit('', 'Работа программы завершена.')

    if len(film_codes) < 1:
        pause_and_exit('В списке 0 фильмов. Работа программы завершена.')
    else:
        with open('./list.txt', 'w') as f:
            f.write('\n'.join(film_codes))
        console.print('Файл "list.txt" сохранен.')

console.print('')
console.print('Начало создания списка.')

err = 0
fullfilmslist = []
for i in range(len(film_codes)):
    if i > 20:
        console.print('[red]Ошибка! Достигнуто ограничение API - не больше 20 фильмов за раз.')
        break
    try:
        filminfo = get_film_info(film_codes[i], api)
        fullfilmslist.append(filminfo)
    except:
        console.print(f'[bold]{film_codes[i]} - ошибка[/bold]')
        err += 1
    else:
        continue

tablenum = len(fullfilmslist)
if tablenum > 1:
    clone_first_table(doc, tablenum - 1)  # добавляем копии шаблонов таблиц
elif tablenum < 1:
    pause_and_exit('[red]Ошибка! Нет фильмов для записи. Список не создан.', '', 'Работа программы завершена.')

# запись информации в таблицы
for i in range(tablenum):
    current_table = doc.tables[i]
    write_film_to_table(current_table, fullfilmslist[i])
    console.print(f'{fullfilmslist[i][0]} - [bright_green]ок', highlight=False)

try:
    doc.save('./list.docx')
except PermissionError:
    pause_and_exit('[red]Ошибка! Нет доступа к файлу "list.docx". Список не создан.', '', 'Работа программы завершена.')

if err > 0:
    console.print(f'[red]Выполнено с ошибками! ({err})')
console.print('')
console.print('Список создан.')

console.rule(style='white')

mp4files = glob.glob('*.mp4')
if len(mp4files) < 1:
    console.print('Файлы mp4 не найдены.')
    clean_and_exit()

# запись тегов
console.print('Найдены файлы mp4:')
for file in mp4files:
    console.print(f'"{file}"')
print('')
ask = str(console.input('Начать запись тегов? [white bold](y,1/n,2) ')).lower()
if ask in ("y", "1"):
    for film in fullfilmslist:
        write_tags_to_mp4(film)
    console.print('')
    console.print('Запись тегов завершена.')
elif ask in ('', 'n', '2'):
    console.print('Отмена. Работа программы завершена.')

clean_and_exit()
